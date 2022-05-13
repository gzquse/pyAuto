from docx import Document #创建文档
from docx.oxml.ns import qn #中文
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #段落
from docx.shared import Pt,RGBColor,Mm,Cm#大小磅数/字号

#下面进入到 模板制作的部分 word制作
#创建模板 word模板 并保存
def batch_pg_add(content,paragraph,space=True):
    for index,text in enumerate(content): #注意这里使用枚举方法获取一下顺序号
        if space == True:
            p = paragraph.add_run('   '+str(index+1)+'.'+text+'\n') #特别注意这里！
            p.font.size = Pt(10)
        else:
            p = paragraph.add_run(str(index + 1) + '.' + text + '\n')  # 特别注意这里！
            p.font.size = Pt(10)

def letter_template(parent_name,parent_title,student_name,score,rank):
    '''
    :param parent_name: 父母姓名
    :param parent_title: 爸爸/妈妈
    :param student_name: 学生姓名
    :param score:  得分
    :param rank: 排名
    :return:
    '''
    #doc
    word_document = Document()  # 创建word文档对象
    word_document.styles['Normal'].font.name = u'黑体'  # 正文/标题1/标题2 （英文）
    word_document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')  # 中文

    # ----段落创建P1-----
    p1 = word_document.add_paragraph()  # 向word添加段落
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中对齐
    # ----添加文字----
    run_text_1 = p1.add_run('给家长的一封信')
    run_text_1.font.size = Pt(18)
    run_text_1.font.bold = True

    # ----段落创建P2----- 学生姓名和家长称呼
    p2 = word_document.add_paragraph()  # 向word添加段落
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 段落居中对齐
    # ----添加文字----
    run_text_2 = p2.add_run('尊敬的 %s %s : ' %(student_name,parent_title))
    run_text_2.font.size = Pt(14)
    run_text_2.font.bold = True

    # ----段落创建P3----- 开场文字
    p3 = word_document.add_paragraph()  # 向word添加段落
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 段落居中对齐
    # ----添加文字----
    run_text_3 = p3.add_run('    你好！秋风送爽，硕果飘香，经历了疫情、汛情，2020年秋季学期已如期而至。为进一步提高大家的安全意识，让同学们度过一个平安、健康、快乐的新学期，我们邀请各位家长朋友参加新学期家长会，具体安排如下：')
    run_text_3.font.size = Pt(10)

    # ----段落创建P4----- 成绩
    p4 = word_document.add_paragraph()  # 向word添加段落
    # ----添加文字----
    run_text_4_1 = p4.add_run('您孩子开学前测评考试的成绩为：\n')
    run_text_4_1.font.size = Pt(14)
    run_text_4_1.font.bold = True

    run_text_4_2 = p4.add_run('语数外三科总分：%s \n' %score)
    run_text_4_2.font.size = Pt(12)

    run_text_4_3 = p4.add_run('班级排名：%s ' % rank)
    run_text_4_3.font.size = Pt(12)


    # ----段落创建P5 送家长的话
    p5 = word_document.add_paragraph()  # 向word添加段落
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 段落居中对齐
    # ----添加文字----
    run_text_5 = p5.add_run('送家长的话：')
    run_text_5.font.size = Pt(12)
    run_text_5.font.bold = True

    # ----段落创建P6----- 家长的话的内容批量处理
    p6 = word_document.add_paragraph()  # 向word添加段落
    # ----添加文字----
    p1_list = [ #说话的内容
        '永远不要跟教自己孩子的老师过不去，良好有效的沟通才是解决问题的法宝；',
        '要对教过自己孩子的老师心存感激和尊敬，因为你是孩子的榜样，只有你感激和尊敬别人，你的孩子才会感激和尊敬你；',
        '永远不要觉得孩子交给老师，教育就是老师和学校的事，因为教育孩子应该是你终身最重要的事业，家庭教育远比学校教育重要得多；',
        '要想孩子变得优秀，那么请你先优秀起来；',
        '家庭和学校教育配合得越好，孩子的教育才越成功。'
    ]
    batch_pg_add(p1_list,p6)

    # ----段落创建P7 送孩子的话
    p7 = word_document.add_paragraph()  # 向word添加段落
    p7.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 段落居中对齐
    # ----添加文字----
    run_text_7 = p7.add_run('送孩子的话：')
    run_text_7.font.size = Pt(12)
    run_text_7.font.bold = True

    # ----段落创建P8----- 孩子的话的内容批量处理
    p8 = word_document.add_paragraph()  # 向word添加段落
    # ----添加文字----
    p2_list = [
        '做一个心存感恩的人。感谢一路走来帮助过你的人，这样你的路途中才会经常有“贵人”相助；',
        '做一个勤奋上进的人。环境不能改变，唯有改变的是自己的心态，你的上进心永远不会被“偷走”； ',
        '做一个认真用心的人。你可以不聪明，也可以不漂亮，但是你不可以不用心和认真。努力了不一定马上能看到“成果”，但有一天你面对挑战与困难的时候，之前积淀的“成果”就会变成你面对困难的勇气和力量；',
        '做一个对社会有益的人。以后有一天你可能会成为于万人之上的人物，也可能是一位尘埃般的人物，但你都不可以剑走偏门，不要想到不劳而获或者天上掉馅饼的事，要有明辨是非的能力和眼睛。请做一个对社会益的人！',
        '做一个永远对世界保持好奇的人。充满活力，充满希望，充满探索欲！这样你才不会停止学习的脚步，你的人生才会过得很精彩！'
    ]

    batch_pg_add(p2_list, p8, False)

    # ----段落创建P9 时间落款
    p9 = word_document.add_paragraph()  # 向word添加段落
    p9.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 段落居中对齐
    # ----添加文字----
    run_text_9 = p9.add_run('2020.8.22')
    run_text_9.font.size = Pt(8)

    file_save = 'letters/%s.docx' %student_name
    word_document.save(file_save)
    return file_save

# letter_template()











